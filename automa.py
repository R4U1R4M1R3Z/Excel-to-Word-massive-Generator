import os
import pandas as pd
from docx import Document

# Cargar el archivo Excel
excel_file = "XXX.xlsx"
df = pd.read_excel(excel_file)

# Ruta de la plantilla de Word y carpeta de salida para los documentos generados
template_path = "XXX.docx"
output_folder = "XXX/XXXX"

# Iterar a través de los registros en el DataFrame
for index, row in df.iterrows():
    row1 = row["XXX"]
    row2 = row["xxx"]
    row3 = row["xxx"]
    
    # Cargar la plantilla de Word
    doc = Document(template_path)
    
    # Buscar y reemplazar el marcador con el número de bastidor en párrafos
    for paragraph in doc.paragraphs:
        if "variable" in paragraph.text:
            paragraph.text = paragraph.text.replace("variable", row1)
    
    # Buscar y reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "variable" in cell.text:
                    cell.text = cell.text.replace("variable", row1)
    
    # Guardar el documento Word generado
    output_filename = f"{output_folder}\\XXX {row1}.docx"
    doc.save(output_filename)

print("Documentos generados correctamente.")

# Abrir la carpeta con los archivos generados
os.system(f'explorer "{output_folder}"')
